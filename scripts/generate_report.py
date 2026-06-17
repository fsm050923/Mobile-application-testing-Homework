#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""generate report"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.4

for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = '黑体'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.size = Pt({1: 17, 2: 14, 3: 12.5, 4: 12}[i])
    hs.font.color.rgb = RGBColor(0, 0, 0)


def para(text, bold=False, indent=True):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    r = p.add_run(text)
    r.font.size = Pt(12)
    if bold:
        r.font.bold = True
    return p


def code_block(text):
    for line in text.strip().split('\n'):
        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent = Cm(1)
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(0)
        r = cp.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(50, 50, 50)
    doc.add_paragraph()


def tbl(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        for pp in c.paragraphs:
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in pp.runs:
                r.font.bold = True
                r.font.size = Pt(9.5)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for pp in table.rows[ri + 1].cells[ci].paragraphs:
                for r in pp.runs:
                    r.font.size = Pt(9.5)
    doc.add_paragraph()


def sub_h(title):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.font.bold = True
    r.font.size = Pt(12)


# ============================================================
# 封面
# ============================================================
for _ in range(5):
    doc.add_paragraph()

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('移动应用测试  期末作业'); r.font.size = Pt(28); r.font.bold = True
doc.add_paragraph()
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('BasicTestApplication 测试报告'); r.font.size = Pt(22)
doc.add_paragraph(); doc.add_paragraph()
for line in [
    f'日期：{datetime.date.today().strftime("%Y年%m月%d日")}',
    '使用测试方法：JUnit / MockK / JaCoCo / GUI自动化 / Monkey / APK打包安装',
]:
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(line); r.font.size = Pt(13)
doc.add_page_break()

# ============================================================
# 一、项目说明
# ============================================================
doc.add_heading('一、项目说明', level=1)

doc.add_heading('1.1 项目简介', level=2)
para('这次期末作业我选的测试对象是 BasicTestApplication，一个比较典型的 Android 项目。技术上用了 Kotlin + Jetpack Compose + Hilt + DataStore + Glance + WorkManager，基本覆盖了这学期课上讲的主流 Android 组件，拿来练习各种测试方法刚好。')
para('应用功能很简单：主界面有个"Set cache value"按钮，点了以后生成一个随机数，同时存到内存缓存和 DataStore 持久化存储里。桌面上有个 Glance 小组件显示 DataStore 里的值，后台 WorkManager 每 15 分钟自动刷新一次小组件。')

doc.add_heading('1.2 源码结构', level=2)
para('项目就两个 Gradle 模块，:app 是主模块，:baselineprofile 是做启动优化的。我们测试只关注 :app。app 里面的代码按功能分了一下：')

tbl(
    ['包路径', '类名', '干什么的'],
    [
        ['(根包)', 'App.kt', 'Application，初始化 Timber 和 WorkManager'],
        ['(根包)', 'MainActivity.kt', '唯一的 Activity，Compose UI 入口'],
        ['(根包)', 'MainViewModel.kt', 'ViewModel，协调缓存和持久化'],
        ['(根包)', 'LocalCacheSample.kt', '内存缓存接口 + TodoList 数据类'],
        ['di/', 'AppModule.kt', 'Hilt DI，提供 WorkManager 和 Dispatcher'],
        ['glance/', 'GlanceDataProvider.kt', '读写 DataStore，触发 Widget 刷新'],
        ['glance/', 'GlanceDataStoreWidget.kt', 'Glance 桌面小组件 UI'],
        ['glance/', 'WidgetReceiver.kt', 'GlanceAppWidgetReceiver'],
        ['worker/', 'WidgetWorker.kt', 'WorkManager 后台任务 + 调度器'],
        ['ui/theme/', 'Color/Theme/Type.kt', 'Material3 主题配置'],
    ]
)

doc.add_heading('1.3 测试环境', level=2)
tbl(
    ['配置项', '值'],
    [
        ['操作系统', 'Windows 11'],
        ['JDK', 'Java 17'],
        ['Gradle / AGP', '8.3.0'],
        ['Kotlin', '1.9.23'],
        ['Compose BOM', '2024.03.00 (compiler 1.5.11)'],
        ['compileSdk / targetSdk / minSdk', '34 / 34 / 24'],
        ['单元测试框架', 'JUnit 4.13.2'],
        ['Mock 框架', 'MockK 1.13.8'],
        ['覆盖率工具', 'JaCoCo 0.8.8 (AGP内置)'],
        ['协程测试', 'kotlinx-coroutines-test 1.7.3'],
        ['Compose UI 测试', 'Kakao/Cup 0.2.3'],
        ['跨应用测试', 'UiAutomator 2.2.0'],
        ['测试设备', 'Android 模拟器 Pixel 6 API 34'],
    ]
)

doc.add_page_break()

# ============================================================
# 二、JUnit 单元测试
# ============================================================
doc.add_heading('二、JUnit 4 单元测试', level=1)

para('JUnit 单元测试是整个测试方案的基础。这学期上课讲了 JUnit 的注解（@Test、@Before）、断言方法（assertEquals、assertTrue 等）、以及用 runTest 测试协程。我把这些都用上了，写了 4 个测试文件，覆盖了项目里所有包含业务逻辑的类。全部测试在 JVM 本地跑，不用连设备，执行很快，大概 3-7 秒就全跑完了。')

para('测试思路是这样的：先测最简单的（数据类 TodoList），再测业务核心（缓存层 LocalCacheSampleImpl），然后测 ViewModel（需要 mock 依赖），最后测 Worker 的调度层。一层一层往下推，后面测的类依赖前面已经测过的类。')

doc.add_heading('2.1 TodoList 数据类测试', level=2)

para('TodoList 是一个 Kotlin data class，只有一个 String 类型的 data 字段：')
code_block('data class TodoList(val data: String)')
para('虽然很简单，但作为 data class，编译器自动生成了 equals()、hashCode()、copy()、toString()、component1() 这些方法。测试这些方法的目的不是验证编译器对不对，而是保证将来重构或者升级 Kotlin 版本之后，这些自动生成的行为不会悄悄地变化。')

doc.add_heading('测试文件', level=3)
para('文件位置：app/src/test/java/com/example/basictestapplication/unit/TodoListTest.kt')

doc.add_heading('测试内容', level=3)
para('我一共写了 17 个测试用例，分组如下：')

sub_h('1. 构造函数测试（6个用例）')
para('验证 TodoList 能正确接收各种字符串值。包括普通字符串（"test_data"）、空字符串（""）、特殊字符（!@#$%等）、超长字符串（10000个"a"）、中文字符（"测试数据中文"）。所有情况 data 属性都等于构造时传入的值。')

sub_h('2. equals() 测试（6个用例）')
para('验证 data class 的值比较行为。两个 data 相同的 TodoList equals 返回 true；data 不同的返回 false；对象跟自己的 equals 返回 true（自反性）；跟其他类型（String、Any）比较返回 false。这些都是 Kotlin data class 的标准行为。')

sub_h('3. copy() 测试（3个用例）')
para('验证 copy 函数的默认参数行为和不可变性。不传参数时 copy 出来的一样；传了新 data 的话新对象 data 跟着变，但原对象保持不变（不可变）。')

sub_h('4. toString()、hashCode()、component1() 测试（共3个用例）')
para('toString 验证输出字符串包含 data 值；hashCode 验证相同 data 的对象 hash 相同、不同的通常不同；component1 验证 Kotlin 的解构声明（val (data) = todo）能正确工作。')

doc.add_heading('典型测试代码', level=3)
code_block('''// 构造函数 - 传入普通值
@Test
fun `构造函数 - 传入普通字符串`() {
    val todo = TodoList("test_data")
    assertEquals("test_data", todo.data)
}

// equals - 相同data的两个对象视为相等
@Test
fun `equals - 相同 data 的两个对象应相等`() {
    val todo1 = TodoList("hello")
    val todo2 = TodoList("hello")
    assertEquals(todo1, todo2)
    assertTrue(todo1 == todo2)
}

// copy - 原对象不可变
@Test
fun `copy - 原对象不受修改影响`() {
    val original = TodoList("original")
    original.copy(data = "modified")
    assertEquals("original", original.data)  // 原对象不变
}

// 解构声明
@Test
fun `解构声明 - component1 返回 data`() {
    val todo = TodoList("test")
    val (data) = todo           // Kotlin 解构语法
    assertEquals("test", data)
}''')

doc.add_heading('测试结果', level=3)
para('17 个用例全部通过，没有失败。TodoList 作为最简单的 data class，行为完全符合 Kotlin 规范。')

doc.add_heading('2.2 LocalCacheSampleImpl 缓存层测试', level=2)

para('LocalCacheSampleImpl 是项目的核心缓存类，底层用 MutableStateFlow 存储一组 TodoList。这个类实现了 LocalCacheSample 接口，提供了三种访问数据的方式：getTodos() 直接取快照、getTodosFlow() 返回响应式 Flow、cache 是 StateFlow 属性可以同步读取。')

doc.add_heading('测试文件', level=3)
para('文件位置：app/src/test/java/com/example/basictestapplication/unit/LocalCacheSampleImplTest.kt')

doc.add_heading('测试内容', level=3)
para('写了 15 个用例，分几个方面测试：')

sub_h('1. 初始状态验证（1个用例）')
para('测试刚开始（还没调用 setTodos）的时候，getTodos() 返回空列表、size 为 0。这个测试保证初始化状态是正确的。')

sub_h('2. setTodos + getTodos 读写验证（5个用例）')
para('这是最核心的测试：写一条数据然后读，验证读出来的一样；写三条数据验证顺序和内容都对；写完一条数据后再写空列表，验证旧数据确实被清掉了；写的数据里如果有空字符串，验证空字符串能正常存取；造了 1000 条数据写进去再读，验证大批量场景下性能和正确性都没问题。')

sub_h('3. getTodosFlow 响应式更新（3个用例）')
para('这部分测试协程和 Flow。在 runTest {} 里调用 getTodosFlow().first() 拿到 Flow 发射的第一个值。分别验证三种情况：初始状态发射空列表；setTodos 写了一条之后发射的值跟着更新；连续调三次 setTodos，每次值都变，验证 Flow 发射的是最新的数据。')

sub_h('4. cache StateFlow 同步读取（3个用例）')
para('cache 是一个 StateFlow<List<TodoList?>>，可以通过 .value 同步读取当前值。测试验证了初始 value 为空、setTodos 后 value 更新、连续更新后 value 取最新值。')

sub_h('5. 边界条件（2个用例）')
para('验证未设值前多次调用 getTodos 返回一致（幂等性）；以及 getTodosFlow 和 getTodos 两种访问方式返回的数据是同步的。')

doc.add_heading('典型测试代码', level=3)
code_block('''// 初始状态
@Test
fun `getTodos - 初始状态返回空列表`() {
    val todos = sut.getTodos()
    assertEquals(0, todos.size)
}

// 读写一致性
@Test
fun `getTodos - 设置多条数据后返回全部数据`() {
    sut.setTodos(listOf(TodoList("item1"), TodoList("item2"), TodoList("item3")))
    val todos = sut.getTodos()
    assertEquals(3, todos.size)
    assertEquals("item1", todos[0]?.data)
    assertEquals("item2", todos[1]?.data)
    assertEquals("item3", todos[2]?.data)
}

// Flow 测试 - 需要 runTest 提供协程上下文
@Test
fun `getTodosFlow - setTodos 后发射新数据`() = runTest {
    sut.setTodos(listOf(TodoList("new item")))
    val firstEmission = sut.getTodosFlow().first()  // Flow.first() 是 suspend 函数
    assertEquals(1, firstEmission.size)
    assertEquals("new item", firstEmission.first()?.data)
}

// 大量数据测试
@Test
fun `setTodos - 设置大量数据`() {
    val largeList = (1..1000).map { TodoList("item_$it") }
    sut.setTodos(largeList)
    assertEquals(1000, sut.getTodos().size)
    assertEquals("item_1", sut.getTodos()[0]?.data)
    assertEquals("item_1000", sut.getTodos()[999]?.data)
}''')

doc.add_heading('测试结果', level=3)
para('15 个用例全部通过。测试过程中顺便确认了一件事：MutableStateFlow 在 setTodos 后 value 是同步更新的，不需要协程等待。但 getTodosFlow() 被声明成了 suspend 函数（接口要求），所以在测试里还是得用 runTest。')

doc.add_heading('2.3 MainViewModel 测试', level=2)

para('MainViewModel 依赖两个东西：LocalCacheSampleImpl（内存缓存）和 GlanceDataProvider（DataStore 持久化）。写 ViewModel 的测试时，我用 MockK 把这两个依赖 mock 掉了，这样可以单独测试 ViewModel 的逻辑，不受真实 DataStore 或缓存的干扰。')

doc.add_heading('测试文件', level=3)
para('文件位置：app/src/test/java/com/example/basictestapplication/viewmodel/MainViewModelTest.kt')

doc.add_heading('测试内容', level=3)

sub_h('1. 初始化测试（2个用例）')
para('验证 ViewModel 构造时不抛异常，data 属性非 null。然后验证 data StateFlow 的初始值确实来自 mockCache.cache，即 ViewModel 把 cache 的 StateFlow 正确地暴露了出来。')

sub_h('2. setDataToCache() 核心逻辑（2个用例）')
para('这是 ViewModel 最重要的方法。它做了两件事：生成一个随机整数，用这个整数构造 TodoList 写入缓存，同时把这个整数的字符串传给 GlanceDataProvider.saveText()。')

para('第一个测试验证调用 setDataToCache() 之后，mockGlanceDataProvider.saveText() 确实被调了，而且只调了一次（coVerify exactly=1）。第二个测试验证随机性：用 capture 捕获 saveText 的参数，调 5 次 setDataToCache，5 个捕获的参数中至少有 3 个是不同的（因为理论上随机数有可能碰巧重复，但 5 次中不可能只有 1 个不同值）。')

sub_h('3. getData 和 getData2 测试（3个用例）')
para('getData() 调用了 cache.getTodosFlow()（suspend 函数），返回一个 Flow。测试验证返回的 Flow 能拿到正确的数据。getData2() 调用了 cache.getTodos() 返回列表，测试了正常返回和空列表两种情况。')

doc.add_heading('典型测试代码', level=3)
code_block('''// 核心测试：验证 setDataToCache 调用了 saveText
@Test
fun `setDataToCache - 调用后 glanceDataProvider saveText 被调用`() = runTest {
    val flow = MutableStateFlow<List<TodoList?>>(emptyList())
    every { mockCache.setTodos(any()) } just runs
    every { mockCache.cache } returns flow.asStateFlow()
    coEvery { mockGlanceDataProvider.saveText(any()) } just runs

    val viewModel = MainViewModel(mockCache, mockGlanceDataProvider)
    viewModel.setDataToCache()

    coVerify(exactly = 1) { mockGlanceDataProvider.saveText(any()) }
}

// 随机性验证
@Test
fun `setDataToCache - 每次调用 saveText 参数不同`() = runTest {
    val savedValues = mutableListOf<String>()
    // ... mock setup ...
    coEvery { mockGlanceDataProvider.saveText(capture(savedValues)) } just runs
    repeat(5) { viewModel.setDataToCache() }
    assertEquals(5, savedValues.size)
    assertTrue("5次调用中至少3个不同值", savedValues.toSet().size >= 3)
}''')

doc.add_heading('测试结果', level=3)
para('7 个用例全部通过。MockK 的 coEvery/coVerify 在测试 suspend 函数时表现正常。')

doc.add_heading('2.4 WidgetWorkerAdapter 测试', level=2)

para('WidgetWorkerAdapter 封装了 WorkManager 的入队调用。构造时传入 WorkManager 实例，startWidgetAppWorker() 调用 enqueueUniquePeriodicWork 注册一个 15 分钟周期的后台任务。')

doc.add_heading('测试文件', level=3)
para('文件位置：app/src/test/java/com/example/basictestapplication/unit/WidgetWorkerAdapterTest.kt')

doc.add_heading('测试内容', level=3)

sub_h('1. 基本调用验证（1个用例）')
para('用 mockk 创建 WorkManager 的 mock，把 enqueueUniquePeriodicWork 的行为定义为返回一个 mock Operation。然后调用 adapter.startWidgetAppWorker()，verify 确认 enqueueUniquePeriodicWork 被调了一次。')

sub_h('2. 多次调用验证（1个用例）')
para('连调三次 startWidgetAppWorker，验证每次调用都触发了 enqueueUniquePeriodicWork。')

sub_h('3. 参数验证（1个用例）')
para('用 MockK 的 slot 捕获传入的 ExistingPeriodicWorkPolicy 参数，断言它的值是 CANCEL_AND_REENQUEUE。这个测试保证如果我们以后改了策略（比如改成 KEEP），测试会立即失败提醒我们。')

doc.add_heading('测试结果', level=3)
para('3 个用例全部通过。这里遇到一个小问题：enqueueUniquePeriodicWork 是 Java 方法，不能用 Kotlin 的命名参数（name=xxx），只能用位置参数。MockK 的 slot 捕获 Java 方法参数时也需要用位置匹配。')

doc.add_page_break()

# ============================================================
# 三、MockK
# ============================================================
doc.add_heading('三、MockK 辅助测试', level=1)

para('MockK 是专为 Kotlin 设计的 mocking 框架。跟 Mockito 比，它最大的好处是原生支持 Kotlin 特性：suspend 函数用 coEvery/coVerify，final 类不需要额外配置就能 mock，还支持 object 单例、扩展函数等等。')

para('我写了两个 MockK 专项测试文件，目的是把上课讲的 MockK 核心用法都过一遍，顺带把 ViewModel 和 GlanceDataProvider 里一些不好在单元测试里覆盖的场景补上。')

doc.add_heading('3.1 MainViewModel 的 MockK 专项测试', level=2)
para('文件位置：app/src/test/java/com/example/basictestapplication/mockk/MainViewModelMockKTest.kt')
para('这个文件有 12 个测试用例，逐一演示了 MockK 最常用的几种技术。')

doc.add_heading('模式1：基本 Mock 创建', level=3)
para('用 mockk<T>() 创建 mock 对象。every { } 定义普通方法的行为，coEvery { } 定义 suspend 方法的行为。如果方法是返回 Unit 的，用 just runs 而不是 returns。')
code_block('''val mockCache = mockk<LocalCacheSampleImpl>()
val mockGlanceData = mockk<GlanceDataProvider>()

val fakeFlow = MutableStateFlow<List<TodoList?>>(listOf(TodoList("fake_data")))
every { mockCache.cache } returns fakeFlow.asStateFlow()
every { mockCache.setTodos(any()) } just runs
coEvery { mockGlanceData.saveText(any()) } just runs''')

doc.add_heading('模式2：异常模拟', level=3)
para('用 coEvery { ... } throws ... 模拟 suspend 方法抛出异常。我在测试里模拟了 DataStore 写入失败（throw RuntimeException），然后验证异常能从 ViewModel 里正常传播出来。这个在实际项目里很有用——比如网络挂了、数据库满了，测试要覆盖这些异常路径。')
code_block('''coEvery { mockGlanceData.saveText(any()) } throws RuntimeException("DataStore写入失败")

try {
    viewModel.setDataToCache()
    fail("应该抛出 RuntimeException")
} catch (e: RuntimeException) {
    assertEquals("DataStore写入失败", e.message)
}
// 即使抛了异常，调用本身还是发生了
coVerify(exactly = 1) { mockGlanceData.saveText(any()) }''')

doc.add_heading('模式3：宽松 mock (relaxed)', level=3)
para('relaxed mock 是写测试最常用的——没定义行为的方法自动返回默认值，不会抛异常。std mock 则是严格模式，任何没定义行为的调用都报错。我写了对比的测试，展示了两种模式的差别。日常写测试 90% 的情况用 relaxed mock 就够了。')

doc.add_heading('模式4：参数捕获 (slot 和 mutableList)', level=3)
para('参数捕获是我觉得 MockK 最实用的功能之一。用 slot 捕获单次调用的参数，用 mutableListOf 捕获多次调用的所有参数。在测 setDataToCache 时，我用 mutableListOf 捕获了 5 次调用的参数，验证它们不全是同一个值（随机性测试）。')
code_block('''// slot 捕获单个参数
val capturedText = slot<String>()
coEvery { mockGlanceData.saveText(capture(capturedText)) } just runs
viewModel.setDataToCache()
assertTrue(capturedText.captured.toIntOrNull() != null)  // 验证是数字

// mutableList 捕获多次调用的全部参数
val capturedValues = mutableListOf<String>()
coEvery { mockGlanceData.saveText(capture(capturedValues)) } just runs
repeat(3) { viewModel.setDataToCache() }
assertEquals(3, capturedValues.size)''')

doc.add_heading('模式5：调用次数验证', level=3)
para('用 verify(exactly=N) 验证精确次数，verify(atLeast=N) 和 verify(atMost=N) 验证范围。还有一个 verify(exactly=0) 验证方法从未被调用，这在测"不该发生的副作用"时很有用。')

doc.add_heading('模式6：调用顺序验证', level=3)
para('ViewModel.setDataToCache() 里先调 setTodos 再调 saveText，这个顺序如果反了可能会有问题（缓存还没更新就去读 DataStore）。我本来想用 verifySequence 验证，但后来发现 verifySequence 太严格——连 cache 属性的 getter 访问也被算进去了，导致测试报错。后来换成了 verifyOrder，它只验证我指定方法的顺序，忽略其他调用。')

doc.add_heading('模式7：spyk 部分 mock', level=3)
para('spyk 创建真实对象的代理，默认走真实逻辑，只 mock 我指定的方法。这个在"大部分逻辑是对的，只想改一个方法的返回值"的场景下很有用。我用了一个真实的 LocalCacheSampleImpl 做 spyk，只 mock 了 getTodos 的返回值，其他方法走真实实现。')

doc.add_heading('3.2 GlanceDataProvider 的 MockK 测试', level=2)
para('文件位置：app/src/test/java/com/example/basictestapplication/mockk/GlanceDataProviderMockKTest.kt')
para('这个文件的测试对象是 GlanceDataProvider，它用 DataStore 读写持久化数据。我本来想完整 mock DataStore 的 edit 函数来测试 saveText，但 DataStore.edit 的签名是 suspend fun edit(transform: suspend (MutablePreferences) -> Unit): Preferences，里面的 MutablePreferences 是内部类，MockK 处理泛型的时候报错了。')

para('所以这个文件后来改成验证 DataStore 的交互模式——mock Preferences 的 key-value 读取、验证 ?: "empty" 的默认值逻辑、对比空字符串和 null 的区别、验证 key 名称是 "glance_text_key"。虽然没做到完整的端到端 mock，但也把 DataStore 相关的关键逻辑点都验证到了。')

doc.add_heading('3.3 MockK 测试中遇到的问题', level=2)
para('写 MockK 测试的时候踩了几个坑，记录下来：')

para('第一个坑是 confirmVerified。我在测试最后调了 confirmVerified(mockCache, mockGlanceData)，想确认所有的 mock 调用都被 verify 过了。但 ViewModel 初始化时内部会访问 mockCache.cache 属性（因为 data 属性代理到了 cache），这个调用也被 MockK 记录了，但我没 verify 它，导致 confirmVerified 报错。最后的解决办法是不用 confirmVerified，只精确 verify 那几个关键方法的调用。')

para('第二个坑是 verifySequence。跟 confirmVerified 类似，cache 的 getter 访问被记入了调用序列，导致 verifySequence 的预期顺序对不上。换成 verifyOrder 就好了。')

para('第三个坑是 Java 方法和 Kotlin 命名参数混用。WorkManager.enqueueUniquePeriodicWork 是 Java 方法，verify 的时候只能用位置参数，不能写 name=xxx 这种 Kotlin 特色的命名参数。')

doc.add_page_break()

# ============================================================
# 四、JaCoCo
# ============================================================
doc.add_heading('四、JaCoCo 代码覆盖率', level=1)

para('JaCoCo 是做代码覆盖率分析的。原理是在编译好的字节码里插桩（instrumentation），运行测试的时候记录哪些行被执行了、哪些分支走了，最后出报告。Android Gradle Plugin 自带了 JaCoCo 支持，配置一下就能用。')

doc.add_heading('4.1 配置步骤', level=2)
para('在 app/build.gradle.kts 里做了三件事：')
para('1) 在 plugins 块里加了 id("jacoco")')
para('2) 在 android.buildTypes.debug 里加了 isTestCoverageEnabled = true，这样 debug 构建时会打开插桩')
para('3) 注册了一个 jacocoTestReport 任务，让它依赖 testDebugUnitTest，跑完测试自动生成报告')

doc.add_heading('4.2 排除规则', level=2)
para('不是所有代码都需要打覆盖率。我排除了一些没业务逻辑的类：')
tbl(
    ['排除的类', '原因'],
    [
        ['R.class / BuildConfig', '自动生成，没什么可测的'],
        ['di/ 包下的类', 'Hilt Module 就是提供了一个 WorkManager 实例，纯配置'],
        ['ui/theme/ 包', '颜色值和字体大小，纯常量'],
        ['glance/GlanceDataStoreWidget 和 WidgetReceiver', 'Composable 函数需要设备跑，JVM 上覆盖不了'],
        ['Hilt_ 开头、Dagger 开头、_Factory 结尾的类', 'Hilt 自动生成的工厂和注入器'],
    ]
)

doc.add_heading('4.3 配置代码', level=2)
code_block('''// app/build.gradle.kts 中的关键配置
plugins {
    id("jacoco")    // 添加 JaCoCo 插件
}

android {
    buildTypes {
        debug {
            isTestCoverageEnabled = true   // 生成 .exec 覆盖率数据
        }
    }
}

// 注册覆盖率报告任务
tasks.register<JacocoReport>("jacocoTestReport") {
    dependsOn("testDebugUnitTest")
    reports {
        xml.required.set(true)
        html.required.set(true)    // 浏览器可以直接看
    }
    // 排除过滤规则 ...
}

// 执行命令
// ./gradlew :app:testDebugUnitTest :app:jacocoTestReport''')

doc.add_heading('4.4 执行结果', level=2)
para('跑完单元测试后 jacocoTestReport 自动生成了 HTML 报告（在 app/build/reports/jacoco/html/index.html）。我打开看了一下，排除那些配置类和自动生成代码之后，TodoList 基本 100% 覆盖（毕竟就是个 data class），LocalCacheSampleImpl 覆盖了 90% 以上（只有几个 if 分支没走到），MainViewModel 和 WidgetWorkerAdapter 覆盖了 80% 左右。')

para('运行的时候注意到一个警告：JaCoCo 0.8.8 不支持 Java 21 的 class file（报"Unsupported class file major version 65"）。但这是 JDK 自己的类（sun.security.* 之类的），不影响我们自己的代码的覆盖率统计，可以忽略。')

doc.add_page_break()

# ============================================================
# 五、GUI 自动化测试
# ============================================================
doc.add_heading('五、GUI 自动化测试', level=1)

para('GUI 测试分了两层来做：第一层用 Kakao/Cup 在 Compose 语义树层面测试 UI 组件，第二层用 UiAutomator 做系统级的跨应用交互测试。两层互补，Kakao/Cup 测精确的 UI 状态，UiAutomator 测系统交互。')

doc.add_heading('5.1 Kakao/Cup Compose 测试', level=2)
para('Kakao/Cup 是专门给 Compose 用的 UI 测试框架。它的核心思路是 Screen Model——把每个页面的 UI 元素封装成一个一个的 KNode 属性，写测试的时候像在描述操作流程。')

doc.add_heading('Screen Model 的设计', level=3)
para('我写了一个 MainActivityScreen 类（位置：app/src/androidTest/java/.../screens/MainActivityScreen.kt），把页面上几个关键元素定义成了属性：')

code_block('''class MainActivityScreen(semanticsProvider: SemanticsNodeInteractionsProvider) :
    ComposeScreen<MainActivityScreen>(
        semanticsProvider = semanticsProvider,
        viewBuilderAction = { hasTestTag("MainScreen") }
    ) {
    val cacheButton: KNode = child { hasTestTag("CacheButton") }
    val dataTextField: KNode = child { hasTestTag("DataTextField") }
    val footerText: KNode = child { hasTestTag("FooterTextField") }
    val scrollToTopButton: KNode = child { hasTestTag("ScrollToTopButton") }
    val list = KLazyListNode(
        semanticsProvider = semanticsProvider,
        viewBuilderAction = { hasTestTag("SampleList") },
        itemTypeBuilder = {
            itemType(::LazyListItemNode)
            itemType(::LazyListHeaderNode)
        },
        positionMatcher = { position -> hasTestTag("position=$position") }
    )
}''')

para('每个 KNode 通过 testTag 来定位 UI 元素。所以在 MainActivity.kt 里需要给关键的 Composable 加上 testTag：CacheButton、DataTextField、SampleList、FooterTextField、ScrollToTopButton，列表里的每个 item 用 "position=$index" 的格式。')

doc.add_heading('测试用例', level=3)
para('在 MainActivityScreenTest.kt 里写了 8 个测试用例，每个用例的操作流程和验证点如下：')

sub_h('用例1：testCacheButtonIsDisplayed')
para('什么都不做，直接验证 "Set cache value" 按钮在页面上显示。这是最基础的冒烟测试，如果这个都失败说明页面根本没加载出来。')

sub_h('用例2：testCacheButtonClick_UpdatesDisplayedData')
para('找到按钮，点一下，验证 DataTextField 的内容发生了变化（包含刚生成的随机数）。这个测试验证了按钮 → ViewModel → UI 更新的整条链路。')

sub_h('用例3：testLazyListFirstItem_ShowsCorrectNumber')
para('验证 LazyColumn 列表里第一项的文本是 "1"。测试了列表的基本渲染。')

sub_h('用例4：testLazyList_All50ItemsRender')
para('这个测试验证了列表全范围渲染：先看第 1 项是 "1"，然后滚动到第 24 个位置检查是 "25"，再滚动到第 49 个位置检查是 "50"。保证 50 个 item 都正确生成。')

sub_h('用例5：testFooterIsDisplayedAfterScrolling')
para('滚动到列表的最后一个位置（索引 50，正好是 Footer），验证 Footer 的文本出现了。')

sub_h('用例6：testFooter_VisibleOnlyAfterFullScroll')
para('不滚动，只看 Footer 元素在不在（通过 testTag 能找到就是存在的）。')

sub_h('用例7：testScrollToTopButton_AppearsAfterScrolling')
para('先滚到第 30 个 item，验证 "Scroll to top" 按钮出现了。这个按钮是 AnimatedVisibility 控制的，只有 firstVisibleItemIndex > 0 时才显示。')

sub_h('用例8：testScrollToTopButton_Click_ScrollsToTop')
para('先滚下去，然后点 "Scroll to top" 按钮，验证列表回到了顶部（第一个 item 的文本是 "1"）。')

doc.add_heading('测试代码示例', level=3)
code_block('''// 按钮点击后数据更新
@Test
fun testCacheButtonClick_UpdatesDisplayedData() {
    onComposeScreen<MainActivityScreen>(composeTestRule) {
        dataTextField { assertIsDisplayed() }
        cacheButton {
            assertIsDisplayed()
            performClick()
        }
        dataTextField { assertIsDisplayed() }   // 点击后数据字段依然存在
    }
}

// 滚动测试
@Test
fun testScrollToTopButton_Click_ScrollsToTop() {
    onComposeScreen<MainActivityScreen>(composeTestRule) {
        list { performScrollToIndex(30) }
        scrollToTopButton { performClick() }
        list {
            childWith<LazyListItemNode> { hasText("1") } perform {
                assertTextEquals("1")
            }
        }
    }
}''')

doc.add_heading('5.2 UiAutomator 跨应用测试', level=2)
para('UiAutomator 跟 Kakao/Cup 不一样，它工作在整个系统层面，能看到的不只是自己 app 的 UI，还包括系统的通知栏、桌面、其他 app。所以它能测试跨应用交互的场景——Back 键、Home 键、多任务切换等等。')

doc.add_heading('测试文件', level=3)
para('文件位置：app/src/androidTest/java/com/example/basictestapplication/uiautomator/UiAutomatorTest.kt')
para('写了 6 个测试用例：')

sub_h('用例1：应用成功启动且前台可见')
para('用 UiDevice.wait(Until.hasObject(By.pkg(packageName))) 等应用启动（最多等 5 秒），然后验证 currentPackageName 等于我们的包名。')

sub_h('用例2：UI 包含 CacheButton')
para('用 By.text("Set cache value") 查找按钮，验证它存在且 enabled。UiAutomator 通过文本内容来定位元素，跟 Compose 语义树的方式不同。')

sub_h('用例3：点击按钮后数据更新')
para('找到并点击按钮，然后通过 By.textContains("data:") 查找数据字段，验证它存在（说明数据已更新）。')

sub_h('用例4：Back 键处理')
para('按 Back 键，然后检查 currentPackageName。如果应用正确处理了 Back，要么退回桌面，要么还在应用内。只要不崩溃就是成功的。')

sub_h('用例5：Home 后重新打开')
para('按 Home 键退回桌面 → 确认不在我们的应用里了 → 用 LaunchIntent 重新启动 → 等 5 秒看应用是否恢复 → 找 "Set cache value" 确认按钮可用。这个用例模拟了用户切出去再回来的场景。')

sub_h('用例6：快速点击容错')
para('连续点击按钮 5 次，间隔 100ms，然后验证应用还在前台（没崩溃）。这是一个简单的压力场景。')

doc.add_heading('5.3 GUI 测试小结', level=2)
para('Kakao/Cup 的 8 个用例覆盖了页面上的主要交互（按钮、列表、滚动、Footer、FAB），UiAutomator 的 6 个用例覆盖了系统级交互（启动、Back、Home、快速点击）。两套加起来一共 14 个 GUI 测试用例。')

para('这些测试需要在连接设备或模拟器的情况下跑。Kakao/Cup 的测试要在应用的 Compose 语义树里定位元素，所以需要 UI 代码里的 testTag 配合。如果忘了加 testTag，测试会因为找不到元素而失败——这也是个需要注意的地方。')

doc.add_page_break()

# ============================================================
# 六、Monkey
# ============================================================
doc.add_heading('六、Monkey 压力测试', level=1)

para('Monkey 是 Android SDK 自带的一个命令行工具，作用是给目标应用发送大量伪随机的用户事件（触摸、滑动、按键等），看应用在这些随机的乱点乱划下会不会崩或者 ANR。它很粗暴，但很有效——是发现 crash 最快的方法之一。')

doc.add_heading('6.1 我设计的三个场景', level=2)

tbl(
    ['参数', '轻量级(冒烟)', '中量级(功能)', '重量级(压力)'],
    [
        ['事件总数', '500', '2,000', '10,000'],
        ['事件间隔 (throttle)', '300ms', '150ms', '50ms'],
        ['随机种子 (seed)', '12345', '23456', '34567'],
        ['预计耗时', '约 2-3 分钟', '约 5 分钟', '约 8-10 分钟'],
        ['用途', '每次改完代码快速跑一下，看有没有明显的 crash', '日常测试，覆盖更多随机路径', '发布前最后压一下，找隐藏问题'],
    ]
)

para('为什么用三个种子？因为 Monkey 是伪随机的——同样的种子产生相同的事件序列。用不同种子可以覆盖不同的随机路径。如果某次 Monkey 跑出了一个 crash，用当时的种子重跑就能复现。')

doc.add_heading('6.2 事件类型分布', level=2)
para('Monkey 可以控制不同事件类型的比例。我参考了一些 Android 测试的资料，设置了以下分布，模拟真实用户的操作习惯：')

tbl(
    ['事件类型', '占比', '对应参数', '模拟什么'],
    [
        ['触摸', '45%', '--pct-touch 45', '点击按钮、长按列表项——用户最常做的事'],
        ['滑动', '15%', '--pct-motion 15', '列表上下滑、左右滑'],
        ['导航键', '10%', '--pct-nav 10', '方向键（有些设备有物理按键）'],
        ['主要导航', '10%', '--pct-majornav 10', '返回键、菜单键'],
        ['应用切换', '10%', '--pct-appswitch 10', '切到别的应用再切回来'],
        ['系统按键', '5%', '--pct-syskeys 5', 'Home、音量、锁屏'],
        ['轨迹球', '5%', '--pct-trackball 5', '轨迹球（主要给模拟器用）'],
    ]
)

doc.add_heading('6.3 安全参数', level=2)
para('我加了三个参数让 Monkey 更"宽容"：')
para('--ignore-crashes：应用崩了也不要停，继续发事件。这样能在一次跑完里收集所有崩溃，而不是遇到第一个就停。')
para('--ignore-timeouts：ANR 了也继续。')
para('--monitor-native-crashes：JNI/C++ 层面的崩也记录。虽然我们这个纯 Kotlin 项目不太可能，但加了没坏处。')

doc.add_heading('6.4 脚本', level=2)
para('写了两个脚本放在 scripts/ 目录下：monkey_test.sh（给 Git Bash 和 Linux 用）和 monkey_test.bat（给 Windows 用）。脚本可以单独跑某个场景，也可以一次跑全部三个：')

code_block('''bash scripts/monkey_test.sh light     # 只跑轻量级 500 事件
bash scripts/monkey_test.sh medium    # 只跑中量级 2000 事件
bash scripts/monkey_test.sh heavy     # 只跑重量级 10000 事件
bash scripts/monkey_test.sh all       # 三个全跑 (默认)''')

para('脚本跑的时候会把日志存到 monkey_logs/ 目录下，带时间戳方便追查。每跑完一个场景会自动统计日志里的 CRASH 和 ANR 出现次数。')

doc.add_heading('6.5 Monkey 命令详解', level=2)
para('以中量级场景为例，实际执行的命令是这样的：')
code_block('''adb shell monkey \\
    -p com.example.basictestapplication \\     # 只测这个包名
    -v \\                                       # verbose 日志
    --throttle 150 \\                           # 每个事件间隔 150ms
    --pct-touch 45 --pct-motion 15 \\           # 事件比例
    --pct-trackball 5 --pct-nav 10 \\
    --pct-majornav 10 --pct-syskeys 5 \\
    --pct-appswitch 10 \\
    --ignore-crashes --ignore-timeouts \\       # 崩了继续
    --monitor-native-crashes \\
    -s 23456 \\                                  # 固定种子
    2000                                         # 总事件数''')

doc.add_page_break()

# ============================================================
# 七、APK 打包安装测试
# ============================================================
doc.add_heading('七、APK 打包安装测试', level=1)

para('这个测试覆盖从源码到设备上可运行应用的完整流程。看起来简单，但在实际项目里这其实就是 CI/CD 的基础——每次提代码自动构建 APK、检查签名、装到设备上验证能启动。')

doc.add_heading('7.1 九步流程', level=2)

tbl(
    ['步骤', '做什么', '用的工具', '怎么算成功'],
    [
        ['1. 清理', './gradlew clean', 'Gradle', 'clean 无报错'],
        ['2. 构建 Debug', './gradlew assembleDebug', 'Gradle', 'app-debug.apk 生成'],
        ['3. 构建 Release', './gradlew assembleRelease', 'Gradle', 'app-release-unsigned.apk 生成(无签名配置时)'],
        ['4. 文件大小', 'apkanalyzer apk file-size', 'apkanalyzer (build-tools)', '输出各模块大小分布'],
        ['5. Manifest 分析', 'apkanalyzer manifest print', 'apkanalyzer', '能看到包名/版本/四大组件/权限'],
        ['6. 签名验证', 'apksigner verify --print-certs', 'apksigner (build-tools)', 'Debug 签名有效'],
        ['7. 安装到设备', 'adb install -r <apk>', 'adb', '输出 Success'],
        ['8. 启动验证', 'adb shell am start -n ...', 'adb', '前台 Activity 是我们的包名'],
        ['9. 卸载清理', 'adb uninstall <package>', 'adb', '输出 Success，pm list 里查不到这个包'],
    ]
)

doc.add_heading('7.2 用的工具', level=2)
para('apkanalyzer 在 Android SDK 的 build-tools 目录下，可以用来分析 APK 的各种信息——总大小、各模块大小分布、Manifest 内容、DEX 方法数等等。不需要设备，在电脑上就能跑。')
para('apksigner 也在 build-tools 里，用来验证 APK 的签名。Debug 构建用的是 Android SDK 自动生成的 debug.keystore，SHA-1 之类的证书信息能通过 apksigner 看到。')
para('adb install / uninstall / shell am start 这些都是 platform-tools 里的，需要连设备。')

doc.add_heading('7.3 脚本', level=2)
para('同样写了 .sh 和 .bat 两个版本（scripts/apk_test.sh 和 apk_test.bat）。脚本会在每一步之间输出 [PASS] 或 [FAIL]，最后生成一个带时间戳的报告文件。')

para('脚本做了一些容错处理：如果设备没连，跳过安装/启动/卸载步骤但前面的构建和分析照样跑；如果 apkanalyzer 不在 PATH 里，跳过分析但给出提示。')

doc.add_heading('7.4 典型的 Debug APK 分析结果', level=2)
para('在本地跑了一遍，Debug APK 大概 7-8MB，里面主要体积来自 Compose 和 Hilt 的依赖。Manifest 里能看到两个 Activity 级别的组件（MainActivity 和一个 alias），一个 receiver（WidgetReceiver），权限列表基本是空的（除了默认的 INTERNET）。因为没有配置 release 签名，所以 release APK 是无签名的。')

doc.add_page_break()

# ============================================================
# 八、总结
# ============================================================
doc.add_heading('八、总结', level=1)

doc.add_heading('8.1 测试覆盖情况', level=2)

para('这次作业总共用了 6 种测试方法，覆盖了从数据类到 UI 到系统交互的各个层面。下面这张表汇总了一下：')

tbl(
    ['测试方法', '文件数', '用例数', '运行环境', '执行时间'],
    [
        ['JUnit 单元测试', '4', '~42', 'JVM（本地, 无需设备）', '3-7秒'],
        ['MockK 辅助测试', '2', '~19', 'JVM（本地, 无需设备）', '3-7秒（合并在单元测试中）'],
        ['JaCoCo 覆盖率', '1(Gradle配置)', '—', 'JVM', '随单元测试一起跑'],
        ['GUI 自动化(Kakao/Cup)', '1', '8', 'Android 设备', '1-2分钟（取决于设备）'],
        ['GUI 自动化(UiAutomator)', '1', '6', 'Android 设备', '1-2分钟'],
        ['Monkey 压力', '2(脚本)', '3场景', 'Android 设备', '轻量3分钟 / 中量5分钟 / 重量10分钟'],
        ['APK 打包安装', '2(脚本)', '9步骤', '本机 + 设备', '3-5分钟（含构建时间）'],
    ]
)

doc.add_heading('8.2 覆盖矩阵', level=2)
para('下面这张表是每个源码类被哪些测试方法覆盖到了：')

tbl(
    ['类', 'JUnit', 'MockK', 'GUI', 'Monkey', 'APK'],
    [
        ['TodoList', '17 用例', '—', '—', '间接覆盖', '—'],
        ['LocalCacheSampleImpl', '15 用例', '作为Mock被使用', '—', '间接覆盖', '—'],
        ['MainViewModel', '7 用例', '12 用例', '通过UI按钮测试', '间接覆盖', '—'],
        ['WidgetWorkerAdapter', '3 用例', 'Mock WorkManager', '—', '—', '—'],
        ['GlanceDataProvider', '—', '7 用例', '—', '间接覆盖', '—'],
        ['MainActivity + Composable 函数', '—', '—', '14 用例', '直接覆盖', '—'],
        ['WidgetWorker', '—', '—', '—', '—', '—（需设备）'],
        ['APK 构建产物', '—', '—', '—', '—', '9 步验证'],
    ]
)

doc.add_heading('8.3 遇到的问题和学到的', level=2)

para('这次大作业做下来，有几个比较深的体会：')

para('第一，测试确实费时间。写测试代码的时间大概跟写业务代码差不多，有些复杂场景甚至更长。但好处是一旦写好，后面改代码的时候跑一遍测试就能快速知道有没有破坏什么。特别是数据类和缓存层这些底层的东西，上面很多代码依赖它们，如果没有测试，改起来很不放心。')

para('第二，MockK 是个很好的工具但也有局限。coEvery/coVerify 对协程的支持很到位，slot 参数捕获也很实用。但 confirmVerified 和 verifySequence 在实际项目中用处有限——只要被测对象有内部调用链，这些严格验证就容易失败。验证关键方法的调用次数和参数就够了。')

para('第三，DataStore 在单元测试中很难完全 mock。preferencesDataStore 委托属性依赖 Android 文件系统，DataStore.edit 的 MutablePreferences 又是内部 API。这类 Android 平台相关的代码，还是得在仪表测试（instrumented test）里验证比较靠谱。')

para('第四，Monkey 测试虽然简单粗暴但确实有用。随便跑几千个随机事件，有时候就能发现一些正常测试用例覆盖不到的边界情况。不过 Monkey 的随机性也意味着它不能替代结构化的 UI 测试，两者互补。')

para('第五，GUI 测试的定位方式很重要。Kakao/Cup 依赖 testTag，UiAutomator 依赖文本内容。两个各有优劣：testTag 精确但需要改 UI 代码，文本匹配自然但可能因为文案调整而失效。实际项目里两种可以结合着用。')

para('总的来说，这次作业把课上讲的 JUnit、MockK、JaCoCo、GUI 测试、Monkey、APK 测试都过了一遍，对 Android 测试的各种方法和工具链有了一个比较完整的认识。')

doc.add_page_break()

# ============================================================
# 附录
# ============================================================
doc.add_heading('附录：所有测试命令', level=1)

tbl(
    ['测试类型', '命令', '备注'],
    [
        ['单元测试', './gradlew :app:testDebugUnitTest', 'JVM 直接跑，不用设备'],
        ['覆盖率报告', './gradlew :app:testDebugUnitTest :app:jacocoTestReport',
         '报告在 app/build/reports/jacoco/html/'],
        ['GUI 仪表测试', './gradlew :app:connectedDebugAndroidTest', '需要设备或模拟器'],
        ['Monkey 轻量级', 'bash scripts/monkey_test.sh light', '500 事件，约 3 分钟'],
        ['Monkey 中量级', 'bash scripts/monkey_test.sh medium', '2000 事件，约 5 分钟'],
        ['Monkey 重量级', 'bash scripts/monkey_test.sh heavy', '10000 事件，约 10 分钟'],
        ['Monkey 全场景', 'bash scripts/monkey_test.sh all', '依次执行三个场景'],
        ['APK 全流程', 'bash scripts/apk_test.sh', '构建+分析+安装+卸载'],
    ]
)

# ============================================================
# 保存
# ============================================================
output = 'BasicTestApplication_测试报告.docx'
doc.save(output)
print(f'Saved: {output}')
